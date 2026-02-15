import streamlit as st
from io import BytesIO
import tempfile
import zipfile
import hashlib
from docx import Document

from src.DocumentGeneratorStrategicActivites import StrategicDocGenerator
from src.DocumentGeneratorBudget import BudgetDocGenerator

st.set_page_config(layout="wide")
st.title("Excel → Word Generator")

TEMPLATE_STRATEGIC_ACTIVITIES_PATH = "templates/template_strategic_activities.docx"
TEMPLATE_BUDGET_PATH = "templates/template_budget.docx"


# -------------------- Utility --------------------
def get_file_hash(file_bytes):
    return hashlib.md5(file_bytes).hexdigest()


# -------------------- Cached generation --------------------
@st.cache_data(show_spinner=True)
def cached_generate_documents(file_bytes, generator_class, template_path):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        tmp.write(file_bytes)
        temp_excel_path = tmp.name

    # Instantiate the correct generator class
    generator = generator_class(temp_excel_path, template_path)
    docs_dict = generator.generate_documents()

    # Serialize Word docs to bytes
    serialized_docs = {}
    for division, doc_obj in docs_dict.items():
        buffer = BytesIO()
        doc_obj.save(buffer)
        serialized_docs[division] = buffer.getvalue()

    return serialized_docs


# -------------------- Choose Excel type --------------------
excel_type = st.radio(
    "Select the type of Excel file",
    options=["Strategic Activities", "Budget"]
)

if excel_type == "Strategic Activities":
    generator_class = StrategicDocGenerator
    template_path = TEMPLATE_STRATEGIC_ACTIVITIES_PATH
    default_prefix = "Review of Strategic Activities of"
else:
    generator_class = BudgetDocGenerator
    template_path = TEMPLATE_BUDGET_PATH
    default_prefix = "Review of Budget of"

# -------------------- User can override filename prefix --------------------
filename_prefix = st.text_input(
    "Optional: File name prefix for each division",
    value=default_prefix
)


# -------------------- Upload Excel --------------------
uploaded_excel = st.file_uploader("Upload Excel File", type=["xlsx"])

# Initialize session state
if "docs_bytes" not in st.session_state:
    st.session_state.docs_bytes = None
if "file_hash" not in st.session_state:
    st.session_state.file_hash = None

if uploaded_excel:
    file_bytes = uploaded_excel.read()
    current_hash = get_file_hash(file_bytes)

    if st.button("Generate Documents"):

        # Only regenerate if file changed
        if st.session_state.file_hash != current_hash:

            docs_bytes = cached_generate_documents(
                file_bytes,
                generator_class,
                template_path
            )

            st.session_state.docs_bytes = docs_bytes
            st.session_state.file_hash = current_hash

        st.success("Documents Generated Successfully!")

    # -------------------- Display Generated Docs --------------------
    if st.session_state.docs_bytes:

        docs_bytes = st.session_state.docs_bytes

        # -------- ZIP Download --------
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for division, doc_bytes in docs_bytes.items():
                file_name = f"{filename_prefix} {division}.docx"
                zf.writestr(file_name, doc_bytes)

        zip_buffer.seek(0)

        st.download_button(
            label="⬇ Download All as ZIP",
            data=zip_buffer,
            file_name="generated_documents.zip",
            mime="application/zip"
        )

        st.markdown("---")
        st.subheader("Generated Documents")

        # -------- Individual Document Preview --------
        for division, doc_bytes in docs_bytes.items():
            with st.expander(f"📄 {division}", expanded=False):

                doc_obj = Document(BytesIO(doc_bytes))

                col1, col2 = st.columns([6, 2])
                with col1:
                    st.markdown("### Document Preview")
                with col2:
                    file_name = f"{filename_prefix} {division}.docx"
                    st.download_button(
                        label="⬇ Download",
                        data=doc_bytes,
                        file_name=file_name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_{division}"
                    )

                st.markdown("---")

                # Paragraphs
                for para in doc_obj.paragraphs:
                    if para.text.strip():
                        st.write(para.text)

                # Tables
                for table_index, table in enumerate(doc_obj.tables):
                    st.markdown(f"#### Table {table_index + 1}")
                    table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    if table_data:
                        st.table(table_data)