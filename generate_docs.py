import pandas as pd
import re
from docx import Document
from datetime import datetime
import os

def clean_name(text):
    text = str(text).strip().lower()
    text = text.replace("%", "percentage")
    text = re.sub(r"[^\w\s]", "", text)
    text = re.sub(r"\s+", "_", text)
    return text

def read_strategic_excel(path):
    df = pd.read_excel(path, header=[3, 4])

    level_0 = pd.Series(df.columns.get_level_values(0)).ffill()
    level_1 = pd.Series(df.columns.get_level_values(1)).ffill()

    df.columns = [
        f"{clean_name(a)}_{clean_name(b)}"
        for a, b in zip(level_0, level_1)
    ]

    df = df.loc[:, ~df.columns.str.contains("^unnamed", case=False)]
    df = df.dropna(how="all").reset_index(drop=True)

    df = df.apply(
        lambda col: col.map(lambda x: x.strip() if isinstance(x, str) else x)
    )

    return df

def clean_excel(df):
    return df[df[df.columns[0]] != "Totals"]

def extract_activity_columns(df_columns):
    activity_map = {}

    for col in df_columns:
        if (
            col.startswith("name_of_the_division")
            or col.startswith("no_of_branches")
            or "percentage" in col
        ):
            continue

        base = re.sub(
            r"_(activity_)?(no_identified|completed_upto_\d+)$",
            "",
            col,
        )

        if base not in activity_map:
            activity_map[base] = {}

        if col.endswith("no_identified"):
            activity_map[base]["target"] = col

        elif "completed_upto" in col:
            activity_map[base]["completed"] = col

    return activity_map

def populate_table(doc, row_data, df_columns):
    table = doc.tables[0]

    existing_activities = {}

    for i in range(1, len(table.rows)):
        activity_text = table.rows[i].cells[0].text.strip()
        key = clean_name(activity_text)
        existing_activities[key] = i
    
    activity_groups = extract_activity_columns(df_columns)

    for activity in activity_groups:
        target_col = activity_groups[activity].get("target")
        ach_col = activity_groups[activity].get("completed")

        target = row_data.get(target_col, 0)
        ach = row_data.get(ach_col, 0)

        if activity in existing_activities:
            row_index = existing_activities[activity]
            row_cells = table.rows[row_index].cells
        else:
            new_row = table.add_row()
            row_cells = new_row.cells
            row_cells[0].text = activity.replace("_", " ").title()

        row_cells[1].text = str(target)
        row_cells[2].text = str(ach)
        row_cells[3].text = f"{round((ach/target)*100, 2)}%"

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip().upper() == "X":
                    cell.text = "0"

def generate_documents(excel_path, template_path):
    doc_dict = {}
    df = read_strategic_excel(excel_path)
    df = clean_excel(df)
    columns = df.columns
    date_str = extract_activity_columns(columns)["college_campus"]["completed"].split("_")[-1]
    try:

        date_obj = datetime.strptime(date_str, "%d%m%y").date().strftime("%d-%m-%y")
    except:
        date_obj = datetime.strptime(date_str, "%d%m%Y").date().strftime("%d-%m-%Y")
    
    for _, row in df.iterrows():
        division = row.iloc[0]

        doc = Document(template_path)

        for p in doc.paragraphs:
            if "XXXXXXXXXXXXXX" in p.text:
                p.text = p.text.replace("XXXXXXXXXXXXXX", str(division))
            if "Date:" in p.text:
                p.text = p.text.replace(
                    "Date:",
                    f"Date: {pd.Timestamp.now().strftime('%d.%m.%Y')}"
                )
            if "Re: Review of strategic activities as at xxxxxxxx" in p.text:
                p.text = p.text.replace(
                    "Re: Review of strategic activities as at xxxxxxxx",
                    f"Re: Review of strategic activities as at {date_obj}"
                )

        populate_table(doc, row, columns)
        doc_dict[division] = doc

    return doc_dict
        # os.makedirs("Files", exist_ok=True)
        # doc.save(f"Files/{division}.docx")