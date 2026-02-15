from src.DocumentGenerator import DocumentGeneratorClass
import pandas as pd
import re
from docx import Document
from datetime import datetime


class StrategicDocGenerator(DocumentGeneratorClass):
    """
    Strategic document generator for Strategic Activities Excel files.
    """

    def __init__(self, excel_path, template_path=None):
        default_template = "templates/template_strategic_activities.docx"
        super().__init__(excel_path, template_path or default_template)

    def read_excel(self):
        df = pd.read_excel(self.excel_path, header=[3, 4])

        level_0 = pd.Series(df.columns.get_level_values(0)).ffill()
        level_1 = pd.Series(df.columns.get_level_values(1)).ffill()

        df.columns = [
            f"{self.clean_name(a)}_{self.clean_name(b)}"
            for a, b in zip(level_0, level_1)
        ]

        df = df.loc[:, ~df.columns.str.contains("^unnamed", case=False)]
        df = df.dropna(how="all").reset_index(drop=True)
        df = df.apply(lambda col: col.map(lambda x: x.strip() if isinstance(x, str) else x))

        return df

    def extract_activity_columns(self):
        activity_map = {}
        for col in self.columns:
            if col.startswith("name_of_the_division") or col.startswith("no_of_branches") or "percentage" in col:
                continue

            base = re.sub(r"_(activity_)?(no_identified|completed_upto_\d+)$", "", col)
            if base not in activity_map:
                activity_map[base] = {}

            if col.endswith("no_identified"):
                activity_map[base]["target"] = col
            elif "completed_upto" in col:
                activity_map[base]["completed"] = col

        return activity_map

    def populate_table(self, doc, row_data):
        table = doc.tables[0]

        existing_activities = {
            self.clean_name(row.cells[0].text.strip()): i
            for i, row in enumerate(table.rows[1:], start=1)
        }

        activity_groups = self.extract_activity_columns()

        for activity, cols in activity_groups.items():
            target_col = cols.get("target")
            ach_col = cols.get("completed")

            target = row_data.get(target_col, 0)
            ach = row_data.get(ach_col, 0)

            if activity in existing_activities:
                row_cells = table.rows[existing_activities[activity]].cells
            else:
                row_cells = table.add_row().cells
                row_cells[0].text = activity.replace("_", " ").title()

            row_cells[1].text = str(target)
            row_cells[2].text = str(ach)
            row_cells[3].text = f"{round((ach / target) * 100, 2)}%" if target else "0%"

        # Replace "X" with 0
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if cell.text.strip().upper() == "X":
                        cell.text = "0"

    def update_paragraphs(self, doc, row_data, date_obj):
        division = row_data.iloc[0]
        for p in doc.paragraphs:
            if "XXXXXXXXXXXXXX" in p.text:
                p.text = p.text.replace("XXXXXXXXXXXXXX", str(division))
            if "Date:" in p.text:
                p.text = p.text.replace("Date:", f"Date: {pd.Timestamp.now().strftime('%d-%m-%Y')}")
            if "Re: Review of strategic activities as at xxxxxxxx" in p.text:
                p.text = p.text.replace(
                    "Re: Review of strategic activities as at xxxxxxxx",
                    f"Re: Review of strategic activities as at {date_obj}"
                )

    def get_date_from_columns(self):
        """
        Extract date from the 'completed' columns of the Excel.
        Returns a formatted string '%d-%m-%y' or '%d-%m-%Y'.
        """
        activity_map = self.extract_activity_columns()
        date_str = next(
            (v.get("completed", "").split("_")[-1] for v in activity_map.values() if "completed" in v),
            None
        )
        try:
            return datetime.strptime(date_str, "%d%m%y").strftime("%d-%m-%y")
        except:
            try:
                return datetime.strptime(date_str, "%d%m%Y").strftime("%d-%m-%Y")
            except:
                return "Unknown Date"